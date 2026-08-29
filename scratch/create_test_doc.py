"""Create a minimal test .docx document for validation testing."""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Set basic properties
doc.core_properties.title = "Test Document"
doc.core_properties.author = "Test Author"

# Normal style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Add H1
doc.add_heading("Sample Remediated Document", level=1)

# Add H6 page marker
doc.add_heading("Page 1", level=6)

# Add body text
doc.add_paragraph("This is a sample paragraph in the remediated document.")

# Add H2
doc.add_heading("Chapter 1: Introduction", level=2)

doc.add_paragraph(
    "This chapter introduces the concepts covered in this textbook. "
    "The content has been remediated for accessibility compliance."
)

# Add a list
doc.add_paragraph("First list item", style="List Bullet")
doc.add_paragraph("Second list item", style="List Bullet")
doc.add_paragraph("Third list item", style="List Bullet")

# Add H2
doc.add_heading("Chapter 2: Data Tables", level=2)

# Add a table
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
# Header row
headers = ["Name", "Value", "Description"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
# Mark header row
trpr = table.rows[0]._element.get_or_add_trPr()
tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
trpr.append(tbl_header)
# Data
data = [
    ["Alpha", "1.0", "First value"],
    ["Beta", "2.5", "Second value"],
    ["Gamma", "3.7", "Third value"],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

# Add H6 page marker
doc.add_heading("Page 2", level=6)

# Add H2
doc.add_heading("Chapter 3: Conclusion", level=2)
doc.add_paragraph("This concludes the sample document.")

# Add a hyperlink paragraph with raw URL (should fail)
doc.add_paragraph("Visit https://www.example.com for more information.")

output_path = "scratch/test_document.docx"
doc.save(output_path)
print(f"Test document saved to: {output_path}")
