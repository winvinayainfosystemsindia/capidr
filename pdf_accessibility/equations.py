"""Equation rendering: LaTeX -> MathML -> native Word OMML equations.

Pipeline: latex2mathml converts LaTeX to MathML, then Microsoft's own
MML2OMML.XSL stylesheet (bundled under resources/, see constants.py) converts
that MathML to OMML — the XML Word itself uses for editable equation objects.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

from .constants import FONT_NAME, FONT_SIZE, MML2OMML_XSL_CANDIDATES

try:
    import latex2mathml.converter as latex2mathml_converter
except ImportError:
    latex2mathml_converter = None
    print("[WARNING] latex2mathml not installed. Equations will render as plain text.")
    print("[WARNING] Install it with: pip install latex2mathml")

_omml_transform = None
_omml_transform_loaded = False


def _get_omml_transform():
    """Lazily load and cache the MathML -> OMML XSLT transform."""
    global _omml_transform, _omml_transform_loaded
    if _omml_transform_loaded:
        return _omml_transform
    _omml_transform_loaded = True

    xsl_path = next((p for p in MML2OMML_XSL_CANDIDATES if p.exists()), None)
    if xsl_path is None:
        print("[WARNING] MML2OMML.XSL not found. Equations will render as plain text.")
        return None

    try:
        xslt_doc = etree.parse(str(xsl_path))
        _omml_transform = etree.XSLT(xslt_doc)
    except Exception as e:
        print(f"[WARNING] Could not load MathML->OMML stylesheet: {e}")
        _omml_transform = None

    return _omml_transform


def latex_to_omml(latex_str: str):
    """Convert a LaTeX equation string to an <m:oMath> lxml element.

    Returns None if conversion isn't possible (missing deps/stylesheet, or
    the LaTeX couldn't be parsed) so callers can fall back to plain text.
    """
    latex_str = (latex_str or "").strip()
    if not latex_str or latex2mathml_converter is None:
        return None

    transform = _get_omml_transform()
    if transform is None:
        return None

    try:
        mathml_str = latex2mathml_converter.convert(latex_str)
        mathml_doc = etree.fromstring(mathml_str.encode("utf-8"))
        omml_doc = transform(mathml_doc)
        return omml_doc.getroot()
    except Exception as e:
        print(f"  [WARNING] Failed to convert equation to OMML ('{latex_str}'): {e}")
        return None


def _add_plain_text_equation_fallback(paragraph, latex_str):
    """Fallback when LaTeX->OMML conversion fails: show the raw LaTeX as italic text."""
    run = paragraph.add_run(latex_str)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.italic = True


def insert_inline_equation(paragraph, latex_str):
    """Insert an equation as a native Word math run inline within a paragraph."""
    omml = latex_to_omml(latex_str)
    if omml is not None:
        paragraph._p.append(omml)
    else:
        _add_plain_text_equation_fallback(paragraph, latex_str)


def insert_display_equation(doc, latex_str, label=None):
    """Add a standalone display equation as its own paragraph.

    With no label, the equation is simply centered. With a `label` (e.g. an
    equation number like "(1)"), the equation is centered via a center tab
    stop and the label is pushed to the right margin via a right tab stop —
    the standard textbook layout for numbered equations.
    """
    para = doc.add_paragraph()
    omml = latex_to_omml(latex_str)
    usable_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin

    if label:
        para.paragraph_format.tab_stops.add_tab_stop(usable_width // 2, WD_TAB_ALIGNMENT.CENTER)
        para.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
        para.add_run("\t")
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if omml is not None:
        omath_para = etree.Element(qn("m:oMathPara"))
        omath_para.append(omml)
        para._p.append(omath_para)
    else:
        _add_plain_text_equation_fallback(para, latex_str)

    if label:
        label_run = para.add_run(f"\t{label}")
        label_run.font.name = FONT_NAME
        label_run.font.size = FONT_SIZE

    return para
