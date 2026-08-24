"""Assembles the structured JSON produced by Claude into a .docx file."""

import io
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .constants import FONT_NAME, FONT_SIZE
from .equations import insert_display_equation
from .rich_text import add_formatted_text


def set_cell_font(cell, text, bold=False):
    """Set font properties for a table cell, with hyperlink support."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_formatted_text(paragraph, str(text), bold=bold)


def add_paragraph_with_font(doc, text, style=None, bold=False, italic=False,
                            alignment=None):
    """Add a paragraph with consistent Times New Roman 12pt formatting and active hyperlinks."""
    para = doc.add_paragraph(style=style)
    if alignment:
        para.alignment = alignment

    add_formatted_text(para, text, bold=bold, italic=italic)
    return para


def set_heading_font(paragraph):
    """Ensure heading paragraphs use Times New Roman 12pt."""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_page_break(doc):
    """Add a page break to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_break = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(fld_break)


def add_table(doc, headers, rows, caption=None):
    """Add a formatted table to the document."""
    if caption:
        cap_para = add_paragraph_with_font(doc, caption, bold=True, italic=True)
        cap_para.space_after = Pt(6)

    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if num_cols == 0:
        return

    # Determine total rows (header + data)
    total_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set header row
    if headers:
        hdr_row = table.rows[0]
        for i, header_text in enumerate(headers):
            if i < len(hdr_row.cells):
                set_cell_font(hdr_row.cells[i], header_text, bold=True)
        # Mark header row to repeat on new pages
        trpr = hdr_row._element.get_or_add_trPr()
        tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trpr.append(tbl_header)

    # Fill data rows
    data_start = 1 if headers else 0
    for row_idx, row_data in enumerate(rows):
        actual_row_idx = row_idx + data_start
        if actual_row_idx < len(table.rows):
            row = table.rows[actual_row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    set_cell_font(row.cells[col_idx], cell_text)

    # Add spacing after table
    doc.add_paragraph()


def add_footnote_text(doc, marker, text):
    """Add footnote as formatted text at the bottom with hyperlink support."""
    para = doc.add_paragraph()
    # Superscript marker
    marker_run = para.add_run(marker)
    marker_run.font.name = FONT_NAME
    marker_run.font.size = Pt(10)
    marker_run.font.superscript = True
    # Footnote text
    para.add_run(" ")
    add_formatted_text(para, text)


def clean_list_text_prefix(text: str, ordered: bool) -> str:
    """Clean leading list markers (e.g. '1. ', 'a) ', '(i) ', '• ', '- ') from list item text
    to prevent duplicate prefix numbers/bullets when Word list styles are applied."""
    if not text:
        return text
    text = text.strip()
    if ordered:
        text = re.sub(r'^(?:\([0-9a-zA-Z]+\)|[0-9a-zA-Z]+[\.\)])\s*', '', text)
    else:
        text = re.sub(r'^[•\-\*▪●◦○–—]\s*', '', text)
    return text.strip()


def get_list_abstract_num_id(doc, style_name="List Number", default_id=7):
    """Find the abstractNumId used by a list style in python-docx."""
    try:
        style = doc.styles[style_name]
        pPr = style._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_elem = numPr.find(qn('w:numId'))
                if numId_elem is not None:
                    style_num_id = numId_elem.get(qn('w:val'))
                    numbering = doc.part.numbering_part.numbering_definitions._numbering
                    for num in numbering.findall(qn('w:num')):
                        if num.get(qn('w:numId')) == style_num_id:
                            ab_elem = num.find(qn('w:abstractNumId'))
                            if ab_elem is not None:
                                return int(ab_elem.get(qn('w:val')))
    except Exception:
        pass
    return default_id


def create_numbered_list_instance(doc, start=1):
    """Create a new <w:num> instance that restarts numbering at `start` with lvlOverride."""
    try:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_id = get_list_abstract_num_id(doc, "List Number", 7)
        num_ids = [int(num.get(qn('w:numId'))) for num in numbering.findall(qn('w:num'))]
        new_num_id = max(num_ids) + 1 if num_ids else 100

        num_xml = f'''
        <w:num {nsdecls("w")} w:numId="{new_num_id}">
            <w:abstractNumId w:val="{abstract_id}"/>
            <w:lvlOverride w:ilvl="0">
                <w:startOverride w:val="{start}"/>
            </w:lvlOverride>
        </w:num>
        '''
        numbering.append(parse_xml(num_xml))
        return new_num_id
    except Exception:
        return None


def build_docx(data: dict, output_path: str, extracted_images: dict = None):
    """Build a Word document from the structured JSON data.

    Args:
        data: Structured JSON from Claude.
        output_path: Path to save the .docx file.
        extracted_images: Dict mapping (page_idx, image_idx) -> image_bytes (PNG).
    """
    if extracted_images is None:
        extracted_images = {}
    doc = Document()

    # -----------------------------------------------------------------------
    # Set default font for the document
    # -----------------------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    # Set font for all heading styles used
    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.name = FONT_NAME
            heading_style.font.size = FONT_SIZE
            heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black text

    # Set list styles
    for style_name in ["List Bullet", "List Number"]:
        if style_name in doc.styles:
            list_style = doc.styles[style_name]
            list_style.font.name = FONT_NAME
            list_style.font.size = FONT_SIZE

    # Initialize numbering definitions element
    dummy = doc.add_paragraph("dummy", style="List Number")
    dummy._element.getparent().remove(dummy._element)

    numbering_element = None
    if hasattr(doc.part, "numbering_part") and doc.part.numbering_part is not None:
        try:
            numbering_element = doc.part.numbering_part.numbering_definitions._numbering
        except Exception:
            numbering_element = None

    current_ordered_num_id = None
    current_bullet_num_id = None

    # -----------------------------------------------------------------------
    # Document title (Heading 1)
    # -----------------------------------------------------------------------
    title = data.get("title", "Untitled Document")
    title_para = doc.add_heading(title, level=1)
    set_heading_font(title_para)
    for run in title_para.runs:
        run.font.bold = True

    # -----------------------------------------------------------------------
    # Process pages
    # -----------------------------------------------------------------------
    pages = data.get("pages", [])
    total_pages = len(pages)

    for page_idx, page in enumerate(pages):
        page_label = page.get("page_label", str(page_idx + 1))

        # Add Heading 6 page marker
        page_marker = doc.add_heading(page_label, level=6)
        set_heading_font(page_marker)

        # Collect footnotes for this page
        page_footnotes = []

        # Process each element on the page
        elements = page.get("elements", [])
        for elem in elements:
            elem_type = elem.get("type", "paragraph")

            if elem_type == "heading":
                current_list_type = None
                current_ordered_num_id = None
                level = elem.get("level", 2)
                level = max(1, min(level, 4))  # Clamp to 1-4
                text = elem.get("text", "")
                heading_para = doc.add_heading(text, level=level)
                set_heading_font(heading_para)
                # Apply bold for visual hierarchy (all headings get bold)
                for run in heading_para.runs:
                    run.font.bold = True
                    if level >= 3:
                        run.font.italic = True

            elif elem_type == "paragraph":
                current_list_type = None
                current_ordered_num_id = None
                text = elem.get("text", "")
                if not text.strip():
                    doc.add_paragraph()
                    continue
                bold = elem.get("bold", False)
                italic = elem.get("italic", False)
                add_paragraph_with_font(doc, text, bold=bold, italic=italic)

            elif elem_type == "list_item":
                raw_text = elem.get("text", "")
                ordered = elem.get("ordered", False)
                item_level = elem.get("level", 0)

                # Clean any explicit prefixes (e.g., '1. ', '• ') from the text to prevent double bullets/numbers
                text = clean_list_text_prefix(raw_text, ordered)

                style_name = "List Number" if ordered else "List Bullet"
                para = add_paragraph_with_font(doc, text, style=style_name)

                if ordered:
                    if current_list_type != "ordered" or current_ordered_num_id is None:
                        current_ordered_num_id = create_numbered_list_instance(doc, start=1)
                        current_list_type = "ordered"

                    if current_ordered_num_id is not None:
                        num_pr = para._element.get_or_add_pPr().get_or_add_numPr()
                        num_pr.get_or_add_numId().val = current_ordered_num_id
                        if item_level > 0:
                            num_pr.get_or_add_ilvl().val = item_level
                else:
                    current_list_type = "unordered"
                    current_ordered_num_id = None
                    if item_level > 0:
                        num_pr = para._element.get_or_add_pPr().get_or_add_numPr()
                        num_pr.get_or_add_ilvl().val = item_level

            elif elem_type == "table":
                current_list_type = None
                current_ordered_num_id = None
                headers = elem.get("headers", [])
                rows = elem.get("rows", [])
                caption = elem.get("caption", None)
                add_table(doc, headers, rows, caption)

            elif elem_type == "figure":
                current_list_type = None
                current_ordered_num_id = None
                fig_num = elem.get("figure_number", "")
                caption = elem.get("caption", "")
                alt_text = elem.get("alt_text", "")
                description = elem.get("description", "")
                image_index = elem.get("image_index", 0)

                # Determine caption text
                if caption:
                    cap_text = caption
                else:
                    cap_text = f"Figure {fig_num}" if fig_num else "Figure"

                # Try to insert actual image from extracted images
                image_key = (page_idx, image_index)
                image_inserted = False

                if image_key in extracted_images:
                    try:
                        img_bytes = extracted_images[image_key]
                        img_stream = io.BytesIO(img_bytes)
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        inline_shape = run.add_picture(img_stream, width=Inches(5.5))

                        # Set alt text on the image for accessibility
                        inline = inline_shape._inline
                        # docPr element holds the alt text
                        doc_pr = inline.find(qn('wp:docPr'))
                        if doc_pr is None:
                            doc_pr = inline.find('.//' + qn('wp:docPr'))
                        if doc_pr is not None:
                            effective_alt = alt_text if alt_text else (description if description else cap_text)
                            doc_pr.set('descr', effective_alt)
                            doc_pr.set('title', cap_text)

                        image_inserted = True
                        print(f"  [IMG] Inserted image for page {page_label}, image {image_index}")
                    except Exception as e:
                        print(f"  [WARNING] Failed to insert image (page {page_label}, idx {image_index}): {e}")

                if not image_inserted:
                    # Fallback: add text placeholder
                    fig_para = add_paragraph_with_font(
                        doc,
                        f"[{cap_text} — image not available]",
                        italic=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    )

                # Add caption below image
                cap_para = add_paragraph_with_font(
                    doc,
                    cap_text,
                    bold=True,
                    italic=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )
                cap_para.paragraph_format.space_before = Pt(4)
                cap_para.paragraph_format.space_after = Pt(8)

            elif elem_type == "equation":
                current_list_type = None
                current_ordered_num_id = None
                latex_text = elem.get("text", "")
                label = elem.get("label", None)
                insert_display_equation(doc, latex_text, label=label)

            elif elem_type == "footnote":
                marker = elem.get("marker", "")
                text = elem.get("text", "")
                page_footnotes.append((marker, text))

            elif elem_type == "blockquote":
                text = elem.get("text", "")
                bq_para = add_paragraph_with_font(doc, text, italic=True)
                bq_para.paragraph_format.left_indent = Cm(1.5)

            elif elem_type == "code_block":
                text = elem.get("text", "")
                code_para = add_paragraph_with_font(doc, text)
                code_para.paragraph_format.left_indent = Cm(1)

        # Add page footnotes at the bottom
        if page_footnotes:
            # Add a thin separator line
            sep_para = doc.add_paragraph()
            sep_run = sep_para.add_run("_" * 30)
            sep_run.font.name = FONT_NAME
            sep_run.font.size = Pt(8)

            for marker, fn_text in page_footnotes:
                add_footnote_text(doc, marker, fn_text)

        # Add page break (except after the last page)
        if page_idx < total_pages - 1:
            add_page_break(doc)

    # -----------------------------------------------------------------------
    # Endnotes (if any, as body content)
    # -----------------------------------------------------------------------
    endnotes = data.get("endnotes", [])
    if endnotes:
        add_page_break(doc)
        en_heading = doc.add_heading("Notes", level=2)
        set_heading_font(en_heading)

        for endnote in endnotes:
            marker = endnote.get("marker", "")
            text = endnote.get("text", "")
            en_para = doc.add_paragraph()
            marker_run = en_para.add_run(f"{marker}. ")
            marker_run.font.name = FONT_NAME
            marker_run.font.size = FONT_SIZE
            marker_run.font.superscript = True
            add_formatted_text(en_para, text)

    # -----------------------------------------------------------------------
    # Final pass: ensure ALL paragraphs use Times New Roman 12pt
    # -----------------------------------------------------------------------
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), FONT_NAME)

    # Also fix table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = FONT_SIZE

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(output_path)
    print(f"[SUCCESS] Word document saved to: {output_path}")
