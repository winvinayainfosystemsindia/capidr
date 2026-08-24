"""Post-build structural verification: heading/page-marker/font sanity checks."""

from docx import Document

from .constants import FONT_NAME


def verify_document(doc_path: str):
    """Verify the generated document's structure."""
    doc = Document(doc_path)

    h6_labels = []
    heading_counts = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    table_count = 0
    total_paragraphs = 0

    for para in doc.paragraphs:
        total_paragraphs += 1
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
                heading_counts[level] = heading_counts.get(level, 0) + 1
                if level == 6:
                    h6_labels.append(para.text)
            except ValueError:
                pass

    table_count = len(doc.tables)

    print("\n" + "=" * 60)
    print("DOCUMENT VERIFICATION REPORT")
    print("=" * 60)
    print(f"Total paragraphs: {total_paragraphs}")
    print(f"Tables: {table_count}")
    print(f"\nHeading breakdown:")
    for level, count in sorted(heading_counts.items()):
        if count > 0:
            print(f"  Heading {level}: {count}")

    print(f"\nHeading 6 (page markers): {len(h6_labels)}")
    if h6_labels:
        print(f"  First: {h6_labels[0]}")
        print(f"  Last:  {h6_labels[-1]}")
        print(f"  Sequence: {', '.join(h6_labels[:10])}", end="")
        if len(h6_labels) > 10:
            print(f" ... ({len(h6_labels)} total)")
        else:
            print()

    # Check font consistency
    non_tnr_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.name != FONT_NAME:
                non_tnr_count += 1

    if non_tnr_count == 0:
        print("\n[OK] All text uses Times New Roman")
    else:
        print(f"\n[WARNING] {non_tnr_count} runs use a font other than Times New Roman")

    print("=" * 60)
