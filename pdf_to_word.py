"""
PDF → Accessible Word Document Converter using Claude Opus API

Usage:
    python pdf_to_word.py <input.pdf> [output.docx]

Requirements:
    pip install anthropic python-docx Pillow

Set your API key:
    set ANTHROPIC_API_KEY=sk-ant-...
"""

import sys
import os
import json
import base64
import re
import io
import html
import argparse
import tempfile
from pathlib import Path

import anthropic
import docx
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

try:
    import latex2mathml.converter as latex2mathml_converter
except ImportError:
    latex2mathml_converter = None
    print("[WARNING] latex2mathml not installed. Equations will render as plain text.")
    print("[WARNING] Install it with: pip install latex2mathml")

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
    print("[WARNING] PyMuPDF not installed. Images will NOT be extracted from PDF.")
    print("[WARNING] Install it with: pip install PyMuPDF")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
MODEL = "claude-opus-5"
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
MAX_TOKENS = 128_000  # Claude Opus 5 supports large output

# Microsoft's official MathML -> OMML stylesheet, bundled locally so equation
# rendering doesn't depend on Office being installed on the machine running
# this script.
MML2OMML_XSL_CANDIDATES = [
    Path(__file__).resolve().parent / "resources" / "MML2OMML.XSL",
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
]

# The skill content is embedded as the system prompt
SYSTEM_PROMPT = r"""
You are an expert document remediation assistant. Your task is to convert a PDF
document into a structured JSON representation that will be used to build an
accessible Word (.docx) document.

<skill>
# PDF → Accessible Word Remediation

A repeatable spec for turning a PDF into a Word file that meets WCAG 2.1 AA /
Section 508 / PDF-UA expectations, without altering the substance of the source.

## Rules

### Page-by-page assembly
For each physical PDF page, in order:
1. A Heading 6 paragraph containing the page's label — the printed page number
   if shown, otherwise "Page N" using PDF sequential position.
2. That page's content with proper formatting.
3. A page break — except after the very last page.

### Heading levels
- Heading 1: document title (once).
- Heading 2/3/4: section headings mapped to the source's hierarchy.

### Global formatting
- Font: Times New Roman, 12pt, for ALL text.
- Preserve heading hierarchy through bold/italics/case, not point size.

### Tables and figures
- Tables: real tables, header row marked, NO merged cells — split merged cells
  into repeated plain cells.
- Figures: sequential "Figure N" captions only for images that carry captions
  in the source. Don't invent captions for decorative images.
- Alt text: every meaningful image gets descriptive alt text. Decorative images
  get empty alt.

### Equations
Every mathematical equation or formula MUST be transcribed as LaTeX and
rendered as a real, editable math object — never as plain text, Unicode math
approximations, or a screenshot/image. Standalone/display equations (on
their own line, e.g. numbered formulas) use the "equation" element type.
Equations that appear inline within a sentence, list item, table cell, or
footnote are written inline in that element's "text" field wrapped in single
dollar signs, e.g. "the area is given by $A = \pi r^2$ for a circle."

### Footnotes and endnotes
Match the source's own choice:
- Bottom-of-page notes → footnotes
- Collected end-of-document notes → body content with own heading

### Never fabricate
If content is illegible or unclear, note it rather than guessing.
</skill>

<output_format>
You MUST respond with a single JSON object. Do NOT include any text before or
after the JSON. The JSON structure is:

{
  "title": "Document Title",
  "pages": [
    {
      "page_label": "1",
      "elements": [
        {
          "type": "heading",
          "level": 2,
          "text": "Section Title"
        },
        {
          "type": "paragraph",
          "text": "Body text content...",
          "bold": false,
          "italic": false
        },
        {
          "type": "list_item",
          "text": "A bullet point",
          "level": 0,
          "ordered": false
        },
        {
          "type": "table",
          "caption": "Table 1: Description",
          "headers": ["Col1", "Col2", "Col3"],
          "rows": [
            ["cell1", "cell2", "cell3"]
          ]
        },
        {
          "type": "figure",
          "figure_number": 1,
          "image_index": 0,
          "caption": "Figure 1: Description of the figure",
          "alt_text": "A detailed, descriptive alt text for screen readers describing exactly what the image shows, its key visual elements, data, and meaning in context",
          "description": "Detailed visual description of the figure"
        },
        {
          "type": "equation",
          "text": "x = \\frac{-b \\pm \\sqrt{b^2 - 4ac}}{2a}",
          "label": "(1)"
        },
        {
          "type": "footnote",
          "marker": "1",
          "text": "Footnote text"
        },
        {
          "type": "blockquote",
          "text": "Quoted text"
        },
        {
          "type": "code_block",
          "text": "code content"
        }
      ]
    }
  ],
  "endnotes": [
    {
      "marker": "1",
      "text": "Endnote text"
    }
  ]
}

Rules for the JSON output:
1. Every page in the PDF MUST have an entry in "pages" with the correct "page_label".
2. page_label should be the printed page number if visible, otherwise "Page N".
3. Heading levels: 1 = doc title, 2/3/4 = content headings. Do NOT use level 5 or 6 in elements (level 6 is reserved for page markers).
4. For tables, split any merged cells into repeated plain cells. Every row must have the same number of columns as headers.
5. For figures with captions in the source, include figure_number in sequence.
6. Include ALL text content — do not skip or summarize anything.
7. For footnotes at the bottom of pages, use type "footnote". For collected endnotes, use the "endnotes" array.
8. Preserve the exact text — this is remediation, not paraphrase.
9. Running headers/footers are NOT content — do not include repeated page furniture.
10. Mid-sentence page breaks: resolve hyphenated words to whole words, split at word boundary.
11. For EVERY image/figure on a page, include a "figure" element with "image_index" set to the 0-based index of that image on its page (first image = 0, second = 1, etc.).
12. The "alt_text" field MUST be a thorough, descriptive text for screen readers. Describe WHAT the image shows in detail — subjects, actions, spatial relationships, colors, text within the image, data values in charts/graphs, and the image's purpose in context. Do NOT use generic descriptions like "An image" or "A figure". Aim for 1-3 sentences that convey the full meaning of the image to someone who cannot see it.
13. If any text contains a web link, URL, or email address (e.g. "https://...", "http://...", "www...", or "mailto:..."), preserve the exact URL or format it as [display text](url) so that it will be rendered as an active, clickable hyperlink in the Word document.
14. LIST FIDELITY (CRITICAL REQUIREMENT):
    - UNORDERED (BULLETED) LISTS: When items in the source PDF begin with a bullet symbol (e.g. •, -, ◦, ▪, ★), you MUST set "ordered": false. NEVER convert bulleted items into numbers (1., 2., etc.).
    - ORDERED (NUMBERED) LISTS: When items in the source PDF begin with numbers or letters (e.g. 1., 2., 3., a., b., (1), (a), i., ii.), you MUST set "ordered": true. NEVER convert numbered items into bullet points.
    - NUMBERING SEQUENCE: Each distinct section or group of items in the PDF is an independent list. Ensure numbered lists restart fresh at 1 for each new section as presented in the PDF.
15. EQUATIONS (CRITICAL REQUIREMENT):
    - Every mathematical equation, formula, or symbolic expression MUST be transcribed as valid LaTeX. Never leave math as plain text, Unicode approximations (e.g. "x² + y²"), or an image placeholder.
    - A standalone/display equation on its own line (commonly numbered) is its own "equation" element, with "text" holding the LaTeX (no surrounding $ delimiters) and, if the source shows an equation number (e.g. "(1)", "(3.2)"), that number goes in "label".
    - An equation that appears inline within a sentence, list item, table cell, or footnote stays inside that element's normal "text" field, wrapped in single dollar signs, e.g. "the identity $e^{i\pi} + 1 = 0$ shows...".
    - Escape backslashes correctly for JSON: a LaTeX command like \frac must appear in the JSON string as \\frac.
    - Transcribe exactly what the source shows (same variables, exponents, subscripts, symbols) — this is remediation, not derivation or simplification.
</output_format>
"""

USER_PROMPT = """Consider the attached PDF and convert it to a structured JSON format
following the output_format specification in your instructions.

Do all these steps precisely:
- Convert ALL content from the PDF — do not skip, summarize, or change any content.
- For each PDF page, create a page entry with the correct page_label.
- Format content headings as Heading 2, 3, and 4 based on the source hierarchy.
- Include and properly format all tables (without merged cells).
- If figures have captions, add figure captions with sequential Figure numbers (Figure 1, Figure 2, etc.).
- Provide appropriate Alt Text for all meaningful images.
- Handle footnotes/endnotes as they appear in the source PDF.
- Transcribe every equation and formula as LaTeX so it renders as a real Word equation, not plain text: standalone/display equations as their own "equation" element, inline equations wrapped in single dollar signs within the surrounding text (e.g. "$A = \\pi r^2$"). Escape backslashes for JSON (\\frac, \\pi, \\sqrt, etc.).
- Identify any URLs or web links in the PDF text and preserve/format them so they become active clickable hyperlinks in the output Word document.
- STRICT LIST FIDELITY: Inspect the PDF carefully for list formatting:
  * If the PDF shows bullet symbols (•, -, ▪), format strictly as UNORDERED bullet list ("ordered": false). Do NOT convert bullets to numbers.
  * If the PDF shows numbers/letters (1., 2., 3., a., b.), format strictly as ORDERED numbered list ("ordered": true).
  * Ensure numbered lists restart at 1 for each new section as in the PDF.
- Ensure page labels are in sequential order.
- Preserve ALL original content verbatim."""


# ---------------------------------------------------------------------------
# PDF Image Extraction
# ---------------------------------------------------------------------------

def extract_images_from_pdf(pdf_path: str) -> dict:
    """Extract images from a PDF, keyed by (page_index, image_index).

    Strategy:
    1. Try extracting embedded raster images per page.
    2. If no embedded images found, render each page as a high-res PNG
       (handles vector graphics, scanned pages, diagrams).

    Returns:
        dict mapping (page_idx, img_idx) -> PNG image bytes
    """
    if fitz is None:
        print("[WARNING] PyMuPDF not available. Skipping image extraction.")
        return {}

    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"[WARNING] PDF not found for image extraction: {pdf_path}")
        return {}

    print("[INFO] Extracting images from PDF...")
    extracted = {}
    try:
        pdf_doc = fitz.open(str(pdf_path))
        num_pages = len(pdf_doc)
        total_images = 0

        # --- Pass 1: Try embedded raster images ---
        for page_idx in range(num_pages):
            page = pdf_doc[page_idx]
            image_list = page.get_images(full=True)
            img_on_page = 0

            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = pdf_doc.extract_image(xref)
                    if base_image is None:
                        continue

                    image_bytes = base_image["image"]
                    image_ext = base_image.get("ext", "png")

                    # Convert to PNG for consistency
                    if image_ext.lower() != "png":
                        try:
                            from PIL import Image as PILImage
                            pil_img = PILImage.open(io.BytesIO(image_bytes))
                            png_buffer = io.BytesIO()
                            pil_img.save(png_buffer, format="PNG")
                            image_bytes = png_buffer.getvalue()
                        except Exception:
                            pass

                    # Skip tiny images (icons, bullets)
                    width = base_image.get("width", 0)
                    height = base_image.get("height", 0)
                    if width < 50 and height < 50:
                        continue

                    extracted[(page_idx, img_on_page)] = image_bytes
                    img_on_page += 1
                    total_images += 1

                except Exception as e:
                    print(f"  [WARNING] Could not extract image from page {page_idx + 1}: {e}")

        # --- Pass 2: If no embedded images, render pages as screenshots ---
        if total_images == 0:
            print("[INFO] No embedded raster images found. Rendering pages as screenshots...")
            for page_idx in range(num_pages):
                page = pdf_doc[page_idx]
                # Render at 200 DPI for good quality
                pix = page.get_pixmap(dpi=200)
                png_bytes = pix.tobytes("png")
                # Store as (page_idx, 0) — one image per page
                extracted[(page_idx, 0)] = png_bytes
                total_images += 1

            print(f"[INFO] Rendered {total_images} page screenshots.")
        else:
            print(f"[INFO] Extracted {total_images} embedded images from {num_pages} pages.")

        pdf_doc.close()

    except Exception as e:
        print(f"[ERROR] Image extraction failed: {e}")

    return extracted


# ---------------------------------------------------------------------------
# Claude API interaction
# ---------------------------------------------------------------------------

def upload_pdf_and_get_response(pdf_path: str, api_key: str) -> dict:
    """Upload PDF to Claude and get structured JSON response."""
    client = anthropic.Anthropic(api_key=api_key, timeout=3600.0)

    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    file_size_mb = pdf_path.stat().st_size / (1024 * 1024)
    print(f"[INFO] PDF file: {pdf_path.name} ({file_size_mb:.1f} MB)")

    # Use Files API for larger files, base64 for smaller ones
    if file_size_mb > 20:
        print("[INFO] Using Files API for large file upload...")
        with open(pdf_path, "rb") as f:
            file_upload = client.files.upload(
                file=(pdf_path.name, f, "application/pdf")
            )
        content = [
            {
                "type": "document",
                "source": {
                    "type": "file",
                    "file_id": file_upload.id,
                },
            },
            {
                "type": "text",
                "text": USER_PROMPT,
            },
        ]
        print(f"[INFO] File uploaded. ID: {file_upload.id}")
    else:
        print("[INFO] Using base64 inline upload...")
        with open(pdf_path, "rb") as f:
            pdf_data = base64.standard_b64encode(f.read()).decode("utf-8")
        content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_data,
                },
            },
            {
                "type": "text",
                "text": USER_PROMPT,
            },
        ]

    print(f"[INFO] Sending request to {MODEL} (streaming enabled)...")
    print("[INFO] Receiving response from Claude...", end="", flush=True)

    response_text = ""
    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": content,
            }
        ],
    ) as stream:
        for text in stream.text_stream:
            response_text += text
            print(".", end="", flush=True)

        message = stream.get_final_message()

    print()
    print(f"[INFO] Response received. Usage: {message.usage}")

    # Check if response was truncated
    if message.stop_reason == "max_tokens":
        print("[WARNING] Response was truncated due to max_tokens limit.")
        print("[WARNING] The output document may be incomplete.")
        print("[WARNING] Consider splitting the PDF into smaller parts.")

    # Parse JSON from response
    json_data = extract_json(response_text)
    return json_data


def extract_json(text: str) -> dict:
    """Extract JSON object from Claude's response text."""
    # Try direct parse first
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try to find JSON within markdown code blocks
    json_match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Try to find first { ... } block
    brace_start = text.find("{")
    if brace_start != -1:
        depth = 0
        for i in range(brace_start, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[brace_start : i + 1])
                    except json.JSONDecodeError:
                        break

    raise ValueError(
        "Could not parse JSON from Claude's response.\n"
        f"Response preview: {text[:500]}..."
    )


# ---------------------------------------------------------------------------
# Equation rendering (LaTeX -> MathML -> native Word OMML equations)
# ---------------------------------------------------------------------------

_omml_transform = None
_omml_transform_loaded = False


def _get_omml_transform():
    """Lazily load and cache the MathML -> OMML XSLT transform."""
    global _omml_transform, _omml_transform_loaded
    if _omml_transform_loaded:
        return _omml_transform
    _omml_transform_loaded = True

    xsl_path = next((p for p in MML2OMML_XSL_CANDIDATES if p.exists()), None)
    if xsl_path is None:
        print("[WARNING] MML2OMML.XSL not found. Equations will render as plain text.")
        return None

    try:
        xslt_doc = etree.parse(str(xsl_path))
        _omml_transform = etree.XSLT(xslt_doc)
    except Exception as e:
        print(f"[WARNING] Could not load MathML->OMML stylesheet: {e}")
        _omml_transform = None

    return _omml_transform


def latex_to_omml(latex_str: str):
    """Convert a LaTeX equation string to an <m:oMath> lxml element.

    Returns None if conversion isn't possible (missing deps/stylesheet, or
    the LaTeX couldn't be parsed) so callers can fall back to plain text.
    """
    latex_str = (latex_str or "").strip()
    if not latex_str or latex2mathml_converter is None:
        return None

    transform = _get_omml_transform()
    if transform is None:
        return None

    try:
        mathml_str = latex2mathml_converter.convert(latex_str)
        mathml_doc = etree.fromstring(mathml_str.encode("utf-8"))
        omml_doc = transform(mathml_doc)
        return omml_doc.getroot()
    except Exception as e:
        print(f"  [WARNING] Failed to convert equation to OMML ('{latex_str}'): {e}")
        return None


def _add_plain_text_equation_fallback(paragraph, latex_str):
    """Fallback when LaTeX->OMML conversion fails: show the raw LaTeX as italic text."""
    run = paragraph.add_run(latex_str)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.italic = True


def insert_inline_equation(paragraph, latex_str):
    """Insert an equation as a native Word math run inline within a paragraph."""
    omml = latex_to_omml(latex_str)
    if omml is not None:
        paragraph._p.append(omml)
    else:
        _add_plain_text_equation_fallback(paragraph, latex_str)


def insert_display_equation(doc, latex_str, label=None):
    """Add a standalone display equation as its own paragraph.

    With no label, the equation is simply centered. With a `label` (e.g. an
    equation number like "(1)"), the equation is centered via a center tab
    stop and the label is pushed to the right margin via a right tab stop —
    the standard textbook layout for numbered equations.
    """
    para = doc.add_paragraph()
    omml = latex_to_omml(latex_str)
    usable_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin

    if label:
        para.paragraph_format.tab_stops.add_tab_stop(usable_width // 2, WD_TAB_ALIGNMENT.CENTER)
        para.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
        para.add_run("\t")
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if omml is not None:
        omath_para = etree.Element(qn("m:oMathPara"))
        omath_para.append(omml)
        para._p.append(omath_para)
    else:
        _add_plain_text_equation_fallback(para, latex_str)

    if label:
        label_run = para.add_run(f"\t{label}")
        label_run.font.name = FONT_NAME
        label_run.font.size = FONT_SIZE

    return para


def add_hyperlink(paragraph, url, text, font_name=FONT_NAME, font_size_pt=12,
                  color="0002D0", underline=True, bold=False, italic=False):
    """Add an active, clickable Word hyperlink to a paragraph."""
    target_url = url
    if not target_url.startswith(("http://", "https://", "mailto:", "ftp://")):
        target_url = "https://" + target_url

    part = paragraph.part
    r_id = part.relate_to(target_url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}/>')
    new_run = parse_xml(f'<w:r {nsdecls("w")}/>')
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')

    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'))
    sz_val = int(font_size_pt * 2)
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{sz_val}"/>'))

    if color:
        rPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>'))
    if underline:
        rPr.append(parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>'))
    if bold:
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    if italic:
        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))

    new_run.append(rPr)
    text_elem = parse_xml(f'<w:t {nsdecls("w")}>{html.escape(text)}</w:t>')
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_formatted_text(paragraph, text, bold=False, italic=False):
    """Add text to a paragraph, automatically detecting markdown links [text](url),
    raw URLs (http://, https://, www., mailto:), and inline LaTeX equations
    ($...$), inserting native Word hyperlinks and equations respectively."""
    if not text:
        return

    # Pattern matches markdown links `[display text](url)`, raw URLs
    # `https://...`, `http://...`, `www....`, `mailto:...`, or inline LaTeX
    # equations delimited by single dollar signs `$...$`.
    pattern = (
        r'(\[(?P<md_text>[^\]]+)\]\((?P<md_url>https?://[^\s)]+|www\.[^\s)]+|mailto:[^\s)]+)\))'
        r'|(?P<raw_url>https?://[^\s)]+|www\.[^\s)]+|mailto:[^\s)]+)'
        r'|\$(?P<equation>[^\s$](?:[^$]*[^\s$])?)\$'
    )

    last_idx = 0
    for match in re.finditer(pattern, text):
        start, end = match.span()
        # Add preceding normal text
        if start > last_idx:
            normal_text = text[last_idx:start]
            run = paragraph.add_run(normal_text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            run.font.bold = bold
            run.font.italic = italic
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)

        equation = match.group("equation")
        if equation:
            insert_inline_equation(paragraph, equation)
            last_idx = end
            continue

        # Handle hyperlink match
        md_text = match.group("md_text")
        md_url = match.group("md_url")
        raw_url = match.group("raw_url")

        if md_text and md_url:
            link_text = md_text
            link_url = md_url
        else:
            link_text = raw_url
            link_url = raw_url

        add_hyperlink(paragraph, link_url, link_text, bold=bold, italic=italic)
        last_idx = end

    # Add remaining text after last match
    if last_idx < len(text):
        remaining_text = text[last_idx:]
        run = paragraph.add_run(remaining_text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = bold
        run.font.italic = italic
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)


def set_cell_font(cell, text, bold=False):
    """Set font properties for a table cell, with hyperlink support."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_formatted_text(paragraph, str(text), bold=bold)


def add_paragraph_with_font(doc, text, style=None, bold=False, italic=False,
                            alignment=None):
    """Add a paragraph with consistent Times New Roman 12pt formatting and active hyperlinks."""
    para = doc.add_paragraph(style=style)
    if alignment:
        para.alignment = alignment

    add_formatted_text(para, text, bold=bold, italic=italic)
    return para


def set_heading_font(paragraph):
    """Ensure heading paragraphs use Times New Roman 12pt."""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_page_break(doc):
    """Add a page break to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_break = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(fld_break)


def add_table(doc, headers, rows, caption=None):
    """Add a formatted table to the document."""
    if caption:
        cap_para = add_paragraph_with_font(doc, caption, bold=True, italic=True)
        cap_para.space_after = Pt(6)

    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if num_cols == 0:
        return

    # Determine total rows (header + data)
    total_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set header row
    if headers:
        hdr_row = table.rows[0]
        for i, header_text in enumerate(headers):
            if i < len(hdr_row.cells):
                set_cell_font(hdr_row.cells[i], header_text, bold=True)
        # Mark header row to repeat on new pages
        trpr = hdr_row._element.get_or_add_trPr()
        tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trpr.append(tbl_header)

    # Fill data rows
    data_start = 1 if headers else 0
    for row_idx, row_data in enumerate(rows):
        actual_row_idx = row_idx + data_start
        if actual_row_idx < len(table.rows):
            row = table.rows[actual_row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    set_cell_font(row.cells[col_idx], cell_text)

    # Add spacing after table
    doc.add_paragraph()


def add_footnote_text(doc, marker, text):
    """Add footnote as formatted text at the bottom with hyperlink support."""
    para = doc.add_paragraph()
    # Superscript marker
    marker_run = para.add_run(marker)
    marker_run.font.name = FONT_NAME
    marker_run.font.size = Pt(10)
    marker_run.font.superscript = True
    # Footnote text
    para.add_run(" ")
    add_formatted_text(para, text)


def clean_list_text_prefix(text: str, ordered: bool) -> str:
    """Clean leading list markers (e.g. '1. ', 'a) ', '(i) ', '• ', '- ') from list item text
    to prevent duplicate prefix numbers/bullets when Word list styles are applied."""
    if not text:
        return text
    text = text.strip()
    if ordered:
        text = re.sub(r'^(?:\([0-9a-zA-Z]+\)|[0-9a-zA-Z]+[\.\)])\s*', '', text)
    else:
        text = re.sub(r'^[•\-\*\u25aa\u25cf\u25e6\u25cb\u2013\u2014]\s*', '', text)
    return text.strip()


def get_list_abstract_num_id(doc, style_name="List Number", default_id=7):
    """Find the abstractNumId used by a list style in python-docx."""
    try:
        style = doc.styles[style_name]
        pPr = style._element.find(docx.oxml.ns.qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(docx.oxml.ns.qn('w:numPr'))
            if numPr is not None:
                numId_elem = numPr.find(docx.oxml.ns.qn('w:numId'))
                if numId_elem is not None:
                    style_num_id = numId_elem.get(docx.oxml.ns.qn('w:val'))
                    numbering = doc.part.numbering_part.numbering_definitions._numbering
                    for num in numbering.findall(docx.oxml.ns.qn('w:num')):
                        if num.get(docx.oxml.ns.qn('w:numId')) == style_num_id:
                            ab_elem = num.find(docx.oxml.ns.qn('w:abstractNumId'))
                            if ab_elem is not None:
                                return int(ab_elem.get(docx.oxml.ns.qn('w:val')))
    except Exception:
        pass
    return default_id


def create_numbered_list_instance(doc, start=1):
    """Create a new <w:num> instance that restarts numbering at `start` with lvlOverride."""
    try:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        abstract_id = get_list_abstract_num_id(doc, "List Number", 7)
        num_ids = [int(num.get(docx.oxml.ns.qn('w:numId'))) for num in numbering.findall(docx.oxml.ns.qn('w:num'))]
        new_num_id = max(num_ids) + 1 if num_ids else 100

        num_xml = f'''
        <w:num {nsdecls("w")} w:numId="{new_num_id}">
            <w:abstractNumId w:val="{abstract_id}"/>
            <w:lvlOverride w:ilvl="0">
                <w:startOverride w:val="{start}"/>
            </w:lvlOverride>
        </w:num>
        '''
        numbering.append(parse_xml(num_xml))
        return new_num_id
    except Exception:
        return None


def build_docx(data: dict, output_path: str, extracted_images: dict = None):
    """Build a Word document from the structured JSON data.
    
    Args:
        data: Structured JSON from Claude.
        output_path: Path to save the .docx file.
        extracted_images: Dict mapping (page_idx, image_idx) -> image_bytes (PNG).
    """
    if extracted_images is None:
        extracted_images = {}
    doc = Document()

    # -----------------------------------------------------------------------
    # Set default font for the document
    # -----------------------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    # Set font for all heading styles used
    for level in range(1, 7):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.name = FONT_NAME
            heading_style.font.size = FONT_SIZE
            heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black text

    # Set list styles
    for style_name in ["List Bullet", "List Number"]:
        if style_name in doc.styles:
            list_style = doc.styles[style_name]
            list_style.font.name = FONT_NAME
            list_style.font.size = FONT_SIZE

    # Initialize numbering definitions element
    dummy = doc.add_paragraph("dummy", style="List Number")
    dummy._element.getparent().remove(dummy._element)

    numbering_element = None
    if hasattr(doc.part, "numbering_part") and doc.part.numbering_part is not None:
        try:
            numbering_element = doc.part.numbering_part.numbering_definitions._numbering
        except Exception:
            numbering_element = None

    current_ordered_num_id = None
    current_bullet_num_id = None

    # -----------------------------------------------------------------------
    # Document title (Heading 1)
    # -----------------------------------------------------------------------
    title = data.get("title", "Untitled Document")
    title_para = doc.add_heading(title, level=1)
    set_heading_font(title_para)
    for run in title_para.runs:
        run.font.bold = True

    # -----------------------------------------------------------------------
    # Process pages
    # -----------------------------------------------------------------------
    pages = data.get("pages", [])
    total_pages = len(pages)

    for page_idx, page in enumerate(pages):
        page_label = page.get("page_label", str(page_idx + 1))

        # Add Heading 6 page marker
        page_marker = doc.add_heading(page_label, level=6)
        set_heading_font(page_marker)

        # Collect footnotes for this page
        page_footnotes = []

        # Process each element on the page
        elements = page.get("elements", [])
        for elem in elements:
            elem_type = elem.get("type", "paragraph")

            if elem_type == "heading":
                current_list_type = None
                current_ordered_num_id = None
                level = elem.get("level", 2)
                level = max(1, min(level, 4))  # Clamp to 1-4
                text = elem.get("text", "")
                heading_para = doc.add_heading(text, level=level)
                set_heading_font(heading_para)
                # Apply bold for visual hierarchy (all headings get bold)
                for run in heading_para.runs:
                    run.font.bold = True
                    if level >= 3:
                        run.font.italic = True

            elif elem_type == "paragraph":
                current_list_type = None
                current_ordered_num_id = None
                text = elem.get("text", "")
                if not text.strip():
                    doc.add_paragraph()
                    continue
                bold = elem.get("bold", False)
                italic = elem.get("italic", False)
                add_paragraph_with_font(doc, text, bold=bold, italic=italic)

            elif elem_type == "list_item":
                raw_text = elem.get("text", "")
                ordered = elem.get("ordered", False)
                item_level = elem.get("level", 0)

                # Clean any explicit prefixes (e.g., '1. ', '• ') from the text to prevent double bullets/numbers
                text = clean_list_text_prefix(raw_text, ordered)

                style_name = "List Number" if ordered else "List Bullet"
                para = add_paragraph_with_font(doc, text, style=style_name)

                if ordered:
                    if current_list_type != "ordered" or current_ordered_num_id is None:
                        current_ordered_num_id = create_numbered_list_instance(doc, start=1)
                        current_list_type = "ordered"

                    if current_ordered_num_id is not None:
                        num_pr = para._element.get_or_add_pPr().get_or_add_numPr()
                        num_pr.get_or_add_numId().val = current_ordered_num_id
                        if item_level > 0:
                            num_pr.get_or_add_ilvl().val = item_level
                else:
                    current_list_type = "unordered"
                    current_ordered_num_id = None
                    if item_level > 0:
                        num_pr = para._element.get_or_add_pPr().get_or_add_numPr()
                        num_pr.get_or_add_ilvl().val = item_level

            elif elem_type == "table":
                current_list_type = None
                current_ordered_num_id = None
                headers = elem.get("headers", [])
                rows = elem.get("rows", [])
                caption = elem.get("caption", None)
                add_table(doc, headers, rows, caption)

            elif elem_type == "figure":
                current_list_type = None
                current_ordered_num_id = None
                fig_num = elem.get("figure_number", "")
                caption = elem.get("caption", "")
                alt_text = elem.get("alt_text", "")
                description = elem.get("description", "")
                image_index = elem.get("image_index", 0)

                # Determine caption text
                if caption:
                    cap_text = caption
                else:
                    cap_text = f"Figure {fig_num}" if fig_num else "Figure"

                # Try to insert actual image from extracted images
                image_key = (page_idx, image_index)
                image_inserted = False

                if image_key in extracted_images:
                    try:
                        img_bytes = extracted_images[image_key]
                        img_stream = io.BytesIO(img_bytes)
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        inline_shape = run.add_picture(img_stream, width=Inches(5.5))

                        # Set alt text on the image for accessibility
                        inline = inline_shape._inline
                        # docPr element holds the alt text
                        doc_pr = inline.find(qn('wp:docPr'))
                        if doc_pr is None:
                            doc_pr = inline.find('.//' + qn('wp:docPr'))
                        if doc_pr is not None:
                            effective_alt = alt_text if alt_text else (description if description else cap_text)
                            doc_pr.set('descr', effective_alt)
                            doc_pr.set('title', cap_text)

                        image_inserted = True
                        print(f"  [IMG] Inserted image for page {page_label}, image {image_index}")
                    except Exception as e:
                        print(f"  [WARNING] Failed to insert image (page {page_label}, idx {image_index}): {e}")

                if not image_inserted:
                    # Fallback: add text placeholder
                    fig_para = add_paragraph_with_font(
                        doc,
                        f"[{cap_text} — image not available]",
                        italic=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    )

                # Add caption below image
                cap_para = add_paragraph_with_font(
                    doc,
                    cap_text,
                    bold=True,
                    italic=True,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )
                cap_para.paragraph_format.space_before = Pt(4)
                cap_para.paragraph_format.space_after = Pt(8)

            elif elem_type == "equation":
                current_list_type = None
                current_ordered_num_id = None
                latex_text = elem.get("text", "")
                label = elem.get("label", None)
                insert_display_equation(doc, latex_text, label=label)

            elif elem_type == "footnote":
                marker = elem.get("marker", "")
                text = elem.get("text", "")
                page_footnotes.append((marker, text))

            elif elem_type == "blockquote":
                text = elem.get("text", "")
                bq_para = add_paragraph_with_font(doc, text, italic=True)
                bq_para.paragraph_format.left_indent = Cm(1.5)

            elif elem_type == "code_block":
                text = elem.get("text", "")
                code_para = add_paragraph_with_font(doc, text)
                code_para.paragraph_format.left_indent = Cm(1)

        # Add page footnotes at the bottom
        if page_footnotes:
            # Add a thin separator line
            sep_para = doc.add_paragraph()
            sep_run = sep_para.add_run("_" * 30)
            sep_run.font.name = FONT_NAME
            sep_run.font.size = Pt(8)

            for marker, fn_text in page_footnotes:
                add_footnote_text(doc, marker, fn_text)

        # Add page break (except after the last page)
        if page_idx < total_pages - 1:
            add_page_break(doc)

    # -----------------------------------------------------------------------
    # Endnotes (if any, as body content)
    # -----------------------------------------------------------------------
    endnotes = data.get("endnotes", [])
    if endnotes:
        add_page_break(doc)
        en_heading = doc.add_heading("Notes", level=2)
        set_heading_font(en_heading)

        for endnote in endnotes:
            marker = endnote.get("marker", "")
            text = endnote.get("text", "")
            en_para = doc.add_paragraph()
            marker_run = en_para.add_run(f"{marker}. ")
            marker_run.font.name = FONT_NAME
            marker_run.font.size = FONT_SIZE
            marker_run.font.superscript = True
            add_formatted_text(en_para, text)

    # -----------------------------------------------------------------------
    # Final pass: ensure ALL paragraphs use Times New Roman 12pt
    # -----------------------------------------------------------------------
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), FONT_NAME)

    # Also fix table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = FONT_SIZE

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(output_path)
    print(f"[SUCCESS] Word document saved to: {output_path}")


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify_document(doc_path: str):
    """Verify the generated document's structure."""
    doc = Document(doc_path)

    h6_labels = []
    heading_counts = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    table_count = 0
    total_paragraphs = 0

    for para in doc.paragraphs:
        total_paragraphs += 1
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
                heading_counts[level] = heading_counts.get(level, 0) + 1
                if level == 6:
                    h6_labels.append(para.text)
            except ValueError:
                pass

    table_count = len(doc.tables)

    print("\n" + "=" * 60)
    print("DOCUMENT VERIFICATION REPORT")
    print("=" * 60)
    print(f"Total paragraphs: {total_paragraphs}")
    print(f"Tables: {table_count}")
    print(f"\nHeading breakdown:")
    for level, count in sorted(heading_counts.items()):
        if count > 0:
            print(f"  Heading {level}: {count}")

    print(f"\nHeading 6 (page markers): {len(h6_labels)}")
    if h6_labels:
        print(f"  First: {h6_labels[0]}")
        print(f"  Last:  {h6_labels[-1]}")
        print(f"  Sequence: {', '.join(h6_labels[:10])}", end="")
        if len(h6_labels) > 10:
            print(f" ... ({len(h6_labels)} total)")
        else:
            print()

    # Check font consistency
    non_tnr_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.name != FONT_NAME:
                non_tnr_count += 1

    if non_tnr_count == 0:
        print("\n✓ All text uses Times New Roman")
    else:
        print(f"\n✗ {non_tnr_count} runs use a font other than Times New Roman")

    print("=" * 60)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def load_env_file():
    """Load variables from .env file if present."""
    env_path = Path(".env")
    if env_path.exists():
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    key = key.strip()
                    value = value.strip().strip("'\"")
                    if key and not os.environ.get(key):
                        os.environ[key] = value


def main():
    load_env_file()
    parser = argparse.ArgumentParser(
        description="Convert PDF to accessible Word document using Claude Opus API"
    )
    parser.add_argument("pdf_path", help="Path to the input PDF file")
    parser.add_argument(
        "output_path",
        nargs="?",
        default=None,
        help="Path for the output .docx file (default: same name as PDF with .docx extension)",
    )
    parser.add_argument(
        "--api-key",
        default=None,
        help="Anthropic API key (or set ANTHROPIC_API_KEY env var)",
    )
    parser.add_argument(
        "--json-only",
        action="store_true",
        help="Only output the JSON structure (don't build .docx)",
    )
    parser.add_argument(
        "--from-json",
        default=None,
        help="Build .docx from a previously saved JSON file (skip API call)",
    )
    parser.add_argument(
        "--save-json",
        action="store_true",
        help="Save the intermediate JSON to a file alongside the .docx",
    )
    parser.add_argument(
        "--verify",
        action="store_true",
        default=True,
        help="Verify the generated document structure (default: True)",
    )

    args = parser.parse_args()

    # Determine output path
    if args.output_path is None:
        args.output_path = str(Path(args.pdf_path).with_suffix(".docx"))

    # Get API key
    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY")

    if args.from_json:
        # Build from existing JSON
        print(f"[INFO] Loading JSON from: {args.from_json}")
        with open(args.from_json, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        if not api_key:
            print("ERROR: No API key provided.")
            print("Set ANTHROPIC_API_KEY environment variable or use --api-key flag.")
            sys.exit(1)

        # Call Claude API
        data = upload_pdf_and_get_response(args.pdf_path, api_key)

    if args.json_only:
        print(json.dumps(data, indent=2, ensure_ascii=False))
        return

    # Save intermediate JSON if requested
    if args.save_json:
        json_path = str(Path(args.output_path).with_suffix(".json"))
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"[INFO] JSON saved to: {json_path}")

    # Extract images from the PDF
    extracted_images = {}
    if not args.from_json:
        extracted_images = extract_images_from_pdf(args.pdf_path)

    # Build the Word document
    print(f"[INFO] Building Word document...")
    build_docx(data, args.output_path, extracted_images)

    # Verify
    if args.verify:
        verify_document(args.output_path)


if __name__ == "__main__":
    main()
