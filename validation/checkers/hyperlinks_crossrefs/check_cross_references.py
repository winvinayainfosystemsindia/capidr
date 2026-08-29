"""Check: Cross references use Word's built-in cross-reference feature.

WCAG 2.2 SC: 1.3.1 Info and Relationships (A), 2.4.5 Multiple Ways (AA)
"""

from typing import List

from docx import Document
from docx.oxml.ns import qn

from validation.base_check import BaseCheck
from validation.models import CheckResult


class CheckCrossReferences(BaseCheck):
    section = "Hyperlinks & Cross References"
    checklist_item = "Cross References"
    description = "Confirm cross references use Word's built-in cross-reference tool"
    wcag_criteria = "1.3.1 Info and Relationships (A), 2.4.5 Multiple Ways (AA)"

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []

        # Look for field codes (REF, PAGEREF) that indicate cross-references
        field_codes = []
        for para in doc.paragraphs:
            for fld_char in para._element.findall(f".//{qn('w:fldChar')}"):
                fld_type = fld_char.get(qn("w:fldCharType"))
                if fld_type == "begin":
                    field_codes.append(True)

        # Check for REF and PAGEREF field instructions
        ref_fields = []
        for para in doc.paragraphs:
            for instr in para._element.findall(f".//{qn('w:instrText')}"):
                if instr.text and ("REF" in instr.text or "PAGEREF" in instr.text):
                    ref_fields.append(instr.text.strip())

        # Also check for internal hyperlinks (anchor-based)
        internal_links = []
        for para in doc.paragraphs:
            for hl in para._element.findall(f".//{qn('w:hyperlink')}"):
                anchor = hl.get(qn("w:anchor"))
                if anchor:
                    internal_links.append(anchor)

        total_refs = len(ref_fields) + len(internal_links)

        if total_refs == 0:
            results.append(
                self.not_applicable(
                    "No cross-references detected in the document"
                )
            )
        else:
            if ref_fields:
                results.append(
                    self.pass_check(
                        location="Entire document",
                        actual=f"{len(ref_fields)} Word cross-reference field(s) found",
                        expected="Cross references should use Word's built-in feature",
                    )
                )
            if internal_links:
                results.append(
                    self.pass_check(
                        location="Entire document",
                        actual=f"{len(internal_links)} internal bookmark link(s) found",
                        expected="Internal links should point to valid bookmarks",
                    )
                )

        return results
