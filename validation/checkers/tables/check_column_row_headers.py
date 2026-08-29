"""Check: Column/row headers are identified appropriately.

WCAG 2.2 SC: 1.3.1 Info and Relationships (A), 4.1.2 Name, Role, Value (A)
"""

from typing import List

from docx import Document

from validation.base_check import BaseCheck
from validation.models import CheckResult
from validation.helpers import get_table_info


class CheckColumnRowHeaders(BaseCheck):
    section = "Tables"
    checklist_item = "Column/Row Headers"
    description = "Check column/row headers are identified appropriately"
    wcag_criteria = "1.3.1 Info and Relationships (A), 4.1.2 Name, Role, Value (A)"

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []

        if not doc.tables:
            results.append(self.not_applicable("No tables found in the document"))
            return results

        for t_idx, table in enumerate(doc.tables):
            info = get_table_info(table)
            table_label = f"Table {t_idx + 1}"

            # Check if first row cells are non-empty (header content exists)
            header_texts = info["header_texts"]
            non_empty_headers = [h for h in header_texts if h]

            if not non_empty_headers:
                results.append(
                    self.fail_check(
                        reason="Table has no identifiable column headers (first row is empty)",
                        location=f"{table_label}, Row 1",
                        expected="First row should contain descriptive column headers",
                        actual="All header cells are empty",
                    )
                )
            elif len(non_empty_headers) < len(header_texts):
                empty_count = len(header_texts) - len(non_empty_headers)
                results.append(
                    self.fail_check(
                        reason=f"{empty_count} header cell(s) are empty",
                        location=f"{table_label}, Row 1",
                        expected="All header cells should have descriptive text",
                        actual=f"Headers: {header_texts}",
                    )
                )
            else:
                # Check if headers are bold (visual indication of header role)
                first_row = table.rows[0]
                all_bold = True
                for cell in first_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip() and not run.font.bold:
                                all_bold = False
                                break

                if all_bold:
                    results.append(
                        self.pass_check(
                            location=f"{table_label}, Row 1",
                            actual=f"Column headers identified and bold: {header_texts}",
                            expected="Headers should be clearly identified",
                        )
                    )
                else:
                    results.append(
                        self.pass_check(
                            location=f"{table_label}, Row 1",
                            actual=f"Column headers present: {header_texts} (not all bold)",
                            expected="Headers should be clearly identified",
                        )
                    )

        return results
