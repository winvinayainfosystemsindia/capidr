"""Check: STEM data tables use accessible Math and notation.

WCAG 2.2 SC: 1.3.1 Info and Relationships (A)
"""

from typing import List

from docx import Document
from docx.oxml.ns import qn

from validation.base_check import BaseCheck
from validation.models import CheckResult


class CheckStemDataTables(BaseCheck):
    section = "Tables"
    checklist_item = "STEM Data Tables"
    description = "Ensure STEM data tables use accessible Math and notation"
    wcag_criteria = "1.3.1 Info and Relationships (A)"

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []

        if not doc.tables:
            results.append(self.not_applicable("No tables found in the document"))
            return results

        # Check tables for equation objects inside cells
        tables_with_math = []
        for t_idx, table in enumerate(doc.tables):
            has_math = False
            has_image_math = False
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Check for OMML equations
                        if para._element.findall(f".//{qn('m:oMath')}"):
                            has_math = True
                        # Check for images (potential equation images)
                        if para._element.findall(f".//{qn('w:drawing')}"):
                            has_image_math = True

            if has_math or has_image_math:
                tables_with_math.append({
                    "index": t_idx,
                    "has_omml": has_math,
                    "has_images": has_image_math,
                })

        if not tables_with_math:
            results.append(
                self.not_applicable(
                    "No STEM data tables with math/equations detected"
                )
            )
        else:
            for tm in tables_with_math:
                table_label = f"Table {tm['index'] + 1}"
                if tm["has_omml"] and not tm["has_images"]:
                    results.append(
                        self.pass_check(
                            location=table_label,
                            actual="Table uses native OMML equations (accessible)",
                            expected="STEM tables should use Word Equation Editor, not images",
                        )
                    )
                elif tm["has_images"]:
                    results.append(
                        self.fail_check(
                            reason="Table contains images that may be equation images instead of native equations",
                            location=table_label,
                            expected="Use Word Equation Editor (OMML) for math in tables",
                            actual="Images found in table cells (possible equation images)",
                        )
                    )

        return results
