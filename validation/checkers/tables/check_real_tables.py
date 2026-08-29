"""Check: Tables are real Word tables (not tab-separated text).

WCAG 2.2 SC: 1.1 Non-text Content (A), 1.3.1 Info and Relationships (A)
"""

import re
from typing import List

from docx import Document

from validation.base_check import BaseCheck
from validation.models import CheckResult
from validation.helpers import get_nearest_page_marker


class CheckRealTables(BaseCheck):
    section = "Tables"
    checklist_item = "Real Tables"
    description = "Convert tables from images to editable Word tables if needed"
    wcag_criteria = "1.1 Non-text Content (A), 1.3.1 Info and Relationships (A)"

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []
        fake_tables = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():
                continue

            # Detect tab-separated data (3+ tabs in a single paragraph)
            tab_count = text.count("\t")
            if tab_count >= 3:
                page = get_nearest_page_marker(doc, idx)
                fake_tables.append({
                    "index": idx,
                    "text": text.strip()[:60],
                    "page": page,
                    "tabs": tab_count,
                })

        if doc.tables:
            results.append(
                self.pass_check(
                    location="Entire document",
                    actual=f"{len(doc.tables)} real Word table(s) found",
                    expected="Data should be in real Word tables, not tab-separated text",
                )
            )

        if fake_tables:
            for ft in fake_tables[:5]:
                results.append(
                    self.fail_check(
                        reason="Tab-separated data detected — should be a real Word table",
                        location=f"Paragraph {ft['index'] + 1} (near {ft['page']}), Text: \"{ft['text']}\"",
                        expected="Use a real Word table for tabular data",
                        actual=f"Paragraph with {ft['tabs']} tabs (fake table)",
                    )
                )
        elif not doc.tables:
            results.append(
                self.not_applicable("No tables or tab-separated data found in the document")
            )

        return results
