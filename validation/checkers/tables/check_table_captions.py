"""Check: Tables have descriptive captions/summaries.

WCAG 2.2 SC: 1.3.1 Info and Relationships (A), 2.4.6 Headings and Labels (AA)
"""

import re
from typing import List

from docx import Document

from validation.base_check import BaseCheck
from validation.models import CheckResult


class CheckTableCaptions(BaseCheck):
    section = "Tables"
    checklist_item = "Table Captions"
    description = "Add descriptive captions/summaries for tables"
    wcag_criteria = "1.3.1 Info and Relationships (A), 2.4.6 Headings and Labels (AA)"

    # Pattern to detect table captions
    _CAPTION_PATTERN = re.compile(r"^table\s+\d+", re.IGNORECASE)

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []

        if not doc.tables:
            results.append(self.not_applicable("No tables found in the document"))
            return results

        # Build a map of paragraph indices to tables (approximate)
        # Find paragraphs that reference "Table N" near each table
        table_elements = [t._element for t in doc.tables]

        for t_idx, table in enumerate(doc.tables):
            table_label = f"Table {t_idx + 1}"
            has_caption = False

            # Search for caption in paragraphs around the table element
            body = doc.element.body
            table_elem = table._element
            prev_sibling = table_elem.getprevious()
            next_sibling = table_elem.getnext()

            # Check previous paragraph
            if prev_sibling is not None and prev_sibling.tag.endswith("}p"):
                text = "".join(
                    t.text or "" for t in prev_sibling.iter()
                    if t.text
                ).strip()
                if self._CAPTION_PATTERN.match(text) or "Caption" in (
                    prev_sibling.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style", "")
                ):
                    has_caption = True

            # Check next paragraph
            if not has_caption and next_sibling is not None and next_sibling.tag.endswith("}p"):
                text = "".join(
                    t.text or "" for t in next_sibling.iter()
                    if t.text
                ).strip()
                if self._CAPTION_PATTERN.match(text):
                    has_caption = True

            if has_caption:
                results.append(
                    self.pass_check(
                        location=table_label,
                        actual="Table has a caption/description",
                        expected="Tables should have descriptive captions",
                    )
                )
            else:
                results.append(
                    self.fail_check(
                        reason="No caption or summary found for this table",
                        location=table_label,
                        expected="Add a descriptive caption (e.g., 'Table 1: Summary of results')",
                        actual="No caption paragraph found adjacent to the table",
                    )
                )

        return results
