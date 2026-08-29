"""Check: Section breaks and page breaks are used correctly.

WCAG 2.2 SC: 1.3.1 Info and Relationships (A), 1.3.2 Meaningful Sequence (A)
"""

from typing import List

from docx import Document
from docx.oxml.ns import qn

from validation.base_check import BaseCheck
from validation.models import CheckResult
from validation.helpers import get_nearest_page_marker


class CheckSectionPageBreaks(BaseCheck):
    section = "Document Setup & Structure"
    checklist_item = "Section & Page Breaks"
    description = "Check section breaks and page breaks are used correctly"
    wcag_criteria = "1.3.1 Info and Relationships (A), 1.3.2 Meaningful Sequence (A)"

    def run(self, doc: Document, doc_path: str) -> List[CheckResult]:
        results = []

        # Count proper page breaks vs multiple empty paragraphs used as spacing
        consecutive_empty = 0
        fake_break_locations = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text and not para.runs:
                consecutive_empty += 1
            else:
                if consecutive_empty >= 5:
                    page = get_nearest_page_marker(doc, idx)
                    fake_break_locations.append(
                        f"Paragraphs {idx - consecutive_empty + 1}-{idx} (near {page})"
                    )
                consecutive_empty = 0

        # Check trailing empty paragraphs
        if consecutive_empty >= 5:
            fake_break_locations.append(
                f"Paragraphs {len(doc.paragraphs) - consecutive_empty + 1}-"
                f"{len(doc.paragraphs)} (end of document)"
            )

        if not fake_break_locations:
            results.append(
                self.pass_check(
                    location="Entire document",
                    actual="Section/page breaks are used properly; no fake breaks detected",
                    expected="Use proper section/page breaks, not multiple empty paragraphs",
                )
            )
        else:
            for loc in fake_break_locations:
                results.append(
                    self.fail_check(
                        reason="Multiple consecutive empty paragraphs used instead of proper page/section break",
                        location=loc,
                        expected="Use Word's built-in page break or section break",
                        actual="5+ consecutive empty paragraphs found (fake spacing)",
                    )
                )

        return results
