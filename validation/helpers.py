"""Shared utility functions used across multiple validation checkers.

Provides consistent helpers for traversing document structure, extracting
elements, and formatting location strings for the validation report.
"""

import re
from typing import Generator, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Paragraph & Run traversal
# ---------------------------------------------------------------------------

def get_all_paragraphs(doc: Document) -> List[Paragraph]:
    """Return all paragraphs in the document body (excluding tables)."""
    return list(doc.paragraphs)


def get_all_paragraphs_with_index(doc: Document) -> Generator[Tuple[int, Paragraph], None, None]:
    """Yield (index, paragraph) for every paragraph in the body."""
    for idx, para in enumerate(doc.paragraphs):
        yield idx, para


def get_table_paragraphs(doc: Document) -> Generator[Tuple[int, int, int, Paragraph], None, None]:
    """Yield (table_idx, row_idx, col_idx, paragraph) for all paragraphs inside tables."""
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    yield t_idx, r_idx, c_idx, para


# ---------------------------------------------------------------------------
# Heading helpers
# ---------------------------------------------------------------------------

def is_heading(para: Paragraph) -> bool:
    """Check if a paragraph uses a built-in Heading style."""
    style_name = para.style.name if para.style else ""
    return style_name.startswith("Heading")


def get_heading_level(para: Paragraph) -> Optional[int]:
    """Extract the heading level from a Heading-styled paragraph. Returns None if not a heading."""
    if not is_heading(para):
        return None
    try:
        return int(para.style.name.split()[-1])
    except (ValueError, IndexError):
        return None


def get_headings(doc: Document) -> List[Tuple[int, int, str]]:
    """Return list of (paragraph_index, heading_level, text) for all headings."""
    headings = []
    for idx, para in enumerate(doc.paragraphs):
        level = get_heading_level(para)
        if level is not None:
            headings.append((idx, level, para.text.strip()))
    return headings


# ---------------------------------------------------------------------------
# Page marker helpers (Heading 6 used as page markers)
# ---------------------------------------------------------------------------

def get_page_markers(doc: Document) -> List[Tuple[int, str]]:
    """Return (paragraph_index, label) for all H6 page-marker paragraphs."""
    markers = []
    for idx, para in enumerate(doc.paragraphs):
        if get_heading_level(para) == 6:
            markers.append((idx, para.text.strip()))
    return markers


def get_nearest_page_marker(doc: Document, para_index: int) -> str:
    """Find the nearest preceding H6 page marker for a given paragraph index."""
    markers = get_page_markers(doc)
    current_page = "Unknown"
    for m_idx, label in markers:
        if m_idx <= para_index:
            current_page = label
        else:
            break
    return current_page


# ---------------------------------------------------------------------------
# Image / inline shape helpers
# ---------------------------------------------------------------------------

def get_images(doc: Document) -> List[dict]:
    """Extract all inline images with their alt text, title, and position.

    Returns a list of dicts with keys:
        - para_index: index of the paragraph containing the image
        - alt_text: the 'descr' attribute from docPr
        - title: the 'title' attribute from docPr
        - name: the 'name' attribute from docPr
        - para_text: surrounding paragraph text
    """
    images = []
    for idx, para in enumerate(doc.paragraphs):
        for drawing in para._element.findall(f".//{qn('w:drawing')}"):
            doc_pr_elements = drawing.findall(f".//{qn('wp:docPr')}")
            for doc_pr in doc_pr_elements:
                images.append({
                    "para_index": idx,
                    "alt_text": doc_pr.get("descr", ""),
                    "title": doc_pr.get("title", ""),
                    "name": doc_pr.get("name", ""),
                    "para_text": para.text.strip()[:80],
                })
    # Also check inside tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for drawing in para._element.findall(f".//{qn('w:drawing')}"):
                        doc_pr_elements = drawing.findall(f".//{qn('wp:docPr')}")
                        for doc_pr in doc_pr_elements:
                            images.append({
                                "para_index": -1,
                                "alt_text": doc_pr.get("descr", ""),
                                "title": doc_pr.get("title", ""),
                                "name": doc_pr.get("name", ""),
                                "para_text": f"Table {t_idx + 1}, Row {r_idx + 1}, Col {c_idx + 1}",
                            })
    return images


# ---------------------------------------------------------------------------
# Equation helpers
# ---------------------------------------------------------------------------

def get_equations(doc: Document) -> List[dict]:
    """Find all OMML equation objects (m:oMath / m:oMathPara) in the document.

    Returns list of dicts with:
        - para_index: paragraph index (-1 if in table)
        - is_display: True if display equation (m:oMathPara), False if inline
        - location: human-readable location string
    """
    equations = []
    math_ns = qn("m:oMath")
    math_para_ns = qn("m:oMathPara")

    for idx, para in enumerate(doc.paragraphs):
        for omath_para in para._element.findall(f".//{math_para_ns}"):
            equations.append({
                "para_index": idx,
                "is_display": True,
                "location": f"Paragraph {idx + 1}",
            })
        for omath in para._element.findall(f"./{math_ns}"):
            # Only count top-level oMath (not those inside oMathPara)
            parent = omath.getparent()
            if parent.tag != math_para_ns:
                equations.append({
                    "para_index": idx,
                    "is_display": False,
                    "location": f"Paragraph {idx + 1}",
                })
    return equations


def has_equations(doc: Document) -> bool:
    """Check if the document contains any OMML equation objects."""
    return len(get_equations(doc)) > 0


# ---------------------------------------------------------------------------
# Hyperlink helpers
# ---------------------------------------------------------------------------

def get_hyperlinks(doc: Document) -> List[dict]:
    """Extract all hyperlinks from the document body.

    Returns list of dicts with:
        - para_index: paragraph index
        - display_text: visible link text
        - url: target URL (or empty if internal)
        - is_external: True if external URL
        - location: human-readable location
    """
    hyperlinks = []
    for idx, para in enumerate(doc.paragraphs):
        for hyperlink_elem in para._element.findall(f".//{qn('w:hyperlink')}"):
            # Get display text
            texts = []
            for run in hyperlink_elem.findall(f".//{qn('w:t')}"):
                if run.text:
                    texts.append(run.text)
            display_text = "".join(texts)

            # Get URL from relationship
            r_id = hyperlink_elem.get(qn("r:id"))
            url = ""
            is_external = False
            if r_id:
                try:
                    rel = para.part.rels.get(r_id)
                    if rel and rel.is_external:
                        url = rel.target_ref
                        is_external = True
                except Exception:
                    pass

            # Check for anchor (internal bookmark link)
            anchor = hyperlink_elem.get(qn("w:anchor")) or hyperlink_elem.get("w:anchor")
            if anchor and not url:
                url = f"#{anchor}"

            hyperlinks.append({
                "para_index": idx,
                "display_text": display_text,
                "url": url,
                "is_external": is_external,
                "location": f"Paragraph {idx + 1}",
            })
    return hyperlinks


# ---------------------------------------------------------------------------
# Footnote / Endnote helpers
# ---------------------------------------------------------------------------

def get_footnote_references(doc: Document) -> List[dict]:
    """Find all footnote reference marks in the document body.

    Returns list of dicts with:
        - para_index: paragraph index
        - footnote_id: the id attribute of the footnote reference
        - location: human-readable location
    """
    refs = []
    for idx, para in enumerate(doc.paragraphs):
        for fn_ref in para._element.findall(f".//{qn('w:footnoteReference')}"):
            fn_id = fn_ref.get(qn("w:id")) or fn_ref.get("w:id") or ""
            refs.append({
                "para_index": idx,
                "footnote_id": fn_id,
                "location": f"Paragraph {idx + 1}",
            })
    return refs


def get_endnote_references(doc: Document) -> List[dict]:
    """Find all endnote reference marks in the document body."""
    refs = []
    for idx, para in enumerate(doc.paragraphs):
        for en_ref in para._element.findall(f".//{qn('w:endnoteReference')}"):
            en_id = en_ref.get(qn("w:id")) or en_ref.get("w:id") or ""
            refs.append({
                "para_index": idx,
                "endnote_id": en_id,
                "location": f"Paragraph {idx + 1}",
            })
    return refs


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def get_table_info(table: Table) -> dict:
    """Extract structured information about a table.

    Returns dict with:
        - row_count, col_count
        - has_header_row: True if first row has tblHeader set
        - merged_cells: list of (row, col) positions with merges
        - header_texts: list of header cell texts
    """
    row_count = len(table.rows)
    col_count = len(table.columns)

    # Check for repeat-header-row (tblHeader)
    has_header_row = False
    if row_count > 0:
        first_row = table.rows[0]
        tr_pr = first_row._element.find(qn("w:trPr"))
        if tr_pr is not None:
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            has_header_row = tbl_header is not None

    # Header texts
    header_texts = []
    if row_count > 0:
        for cell in table.rows[0].cells:
            header_texts.append(cell.text.strip())

    # Merged cells
    merged_cells = []
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tc = cell._element
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                # Horizontal merge (gridSpan)
                grid_span = tc_pr.find(qn("w:gridSpan"))
                if grid_span is not None:
                    span_val = grid_span.get(qn("w:val"), "1")
                    if int(span_val) > 1:
                        merged_cells.append({
                            "row": r_idx,
                            "col": c_idx,
                            "type": "horizontal",
                            "span": int(span_val),
                        })
                # Vertical merge
                v_merge = tc_pr.find(qn("w:vMerge"))
                if v_merge is not None:
                    merge_val = v_merge.get(qn("w:val"), "continue")
                    merged_cells.append({
                        "row": r_idx,
                        "col": c_idx,
                        "type": "vertical",
                        "merge_value": merge_val,
                    })

    return {
        "row_count": row_count,
        "col_count": col_count,
        "has_header_row": has_header_row,
        "merged_cells": merged_cells,
        "header_texts": header_texts,
    }


# ---------------------------------------------------------------------------
# Bookmark helpers
# ---------------------------------------------------------------------------

def get_bookmarks(doc: Document) -> List[dict]:
    """Extract all bookmarks from the document.

    Returns list of dicts with:
        - name: bookmark name
        - id: bookmark id
    """
    bookmarks = []
    for bookmark_start in doc.element.body.findall(f".//{qn('w:bookmarkStart')}"):
        name = bookmark_start.get(qn("w:name")) or bookmark_start.get("w:name", "")
        bm_id = bookmark_start.get(qn("w:id")) or bookmark_start.get("w:id", "")
        if name and not name.startswith("_"):  # Skip internal bookmarks
            bookmarks.append({"name": name, "id": bm_id})
    return bookmarks


# ---------------------------------------------------------------------------
# Text box / floating content helpers
# ---------------------------------------------------------------------------

def get_text_boxes(doc: Document) -> List[dict]:
    """Find text boxes (txbxContent) in the document that may break reading order."""
    text_boxes = []
    for idx, para in enumerate(doc.paragraphs):
        for txbx in para._element.findall(f".//{qn('w:txbxContent')}"):
            texts = []
            for t_para in txbx.findall(f".//{qn('w:t')}"):
                if t_para.text:
                    texts.append(t_para.text)
            text_content = " ".join(texts).strip()
            if text_content:
                text_boxes.append({
                    "para_index": idx,
                    "text": text_content[:100],
                    "location": f"Paragraph {idx + 1}",
                })
    return text_boxes


# ---------------------------------------------------------------------------
# Style detection helpers
# ---------------------------------------------------------------------------

def looks_like_heading(para: Paragraph) -> bool:
    """Heuristically detect if a paragraph *looks* like a heading but doesn't
    use a Heading style (manual bold + larger font)."""
    if is_heading(para):
        return False
    text = para.text.strip()
    if not text or len(text) > 200:
        return False

    # Check if all runs are bold and/or have a larger font
    if not para.runs:
        return False
    all_bold = all(run.font.bold for run in para.runs if run.text.strip())
    has_large_font = any(
        run.font.size and run.font.size > 150000  # > 12pt in EMUs
        for run in para.runs
        if run.text.strip()
    )
    # Short text + all bold + larger font = likely a heading
    if all_bold and len(text) < 100:
        return True
    if has_large_font and len(text) < 100:
        return True
    return False


def is_manual_list(para: Paragraph) -> Optional[str]:
    """Detect if a paragraph starts with a manual list marker instead of using
    Word list styles. Returns the marker string if found, None otherwise."""
    text = para.text.strip()
    if not text:
        return None
    style_name = para.style.name if para.style else ""
    if style_name.startswith("List"):
        return None

    # Check for common manual markers
    patterns = [
        (r"^[•●○◦▪–—]\s", "bullet"),
        (r"^[-]\s", "dash"),
        (r"^[*]\s", "asterisk"),
        (r"^\d+[.)]\s", "numbered"),
        (r"^[a-z][.)]\s", "lettered"),
        (r"^\([a-z0-9]+\)\s", "parenthesized"),
        (r"^[ivxIVX]+[.)]\s", "roman"),
    ]
    for pattern, marker_type in patterns:
        if re.match(pattern, text):
            return marker_type
    return None


# ---------------------------------------------------------------------------
# Location formatting
# ---------------------------------------------------------------------------

def format_location(
    para_index: int = -1,
    page_label: str = "",
    table_index: int = -1,
    row_index: int = -1,
    col_index: int = -1,
    snippet: str = "",
) -> str:
    """Build a human-readable location string for the report."""
    parts = []
    if page_label:
        parts.append(f"Page: {page_label}")
    if para_index >= 0:
        parts.append(f"Paragraph {para_index + 1}")
    if table_index >= 0:
        parts.append(f"Table {table_index + 1}")
    if row_index >= 0:
        parts.append(f"Row {row_index + 1}")
    if col_index >= 0:
        parts.append(f"Col {col_index + 1}")
    if snippet:
        short = snippet[:60] + ("..." if len(snippet) > 60 else "")
        parts.append(f'Text: "{short}"')
    return ", ".join(parts) if parts else "Document level"
